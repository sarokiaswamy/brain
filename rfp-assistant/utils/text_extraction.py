"""
Text extraction utilities for RFP documents.
Specialized for extracting structured content from RFP documents using Azure OpenAI.
"""
import os
import re
import sys
import json
import logging
from typing import Dict, List, Optional, Tuple, Any
import PyPDF2
import docx
from langchain.docstore.document import Document

# Add parent directory to path to import from semantic_similarity
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
from semantic_similarity import get_client, GeminiSettings
from utils.prompt_loader import prompt_loader

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Initialize Gemini settings
gemini_settings = GeminiSettings()

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file with special handling for RFP documents.
    Attempts to preserve structure and formatting relevant to RFPs.
    """
    try:
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
        return text
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {str(e)}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file with special handling for RFP documents.
    Attempts to preserve structure and formatting relevant to RFPs.
    """
    try:
        doc = docx.Document(file_path)
        
        # Extract paragraphs with special handling for headers
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's likely a header (based on style)
                if para.style.name.startswith('Heading'):
                    paragraphs.append(f"\n## {para.text}\n")
                else:
                    paragraphs.append(para.text)
        
        # Extract tables which are common in RFPs
        for table in doc.tables:
            table_text = []
            for i, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                table_text.append(" | ".join(row_text))
            
            # Add table with separator
            paragraphs.append("\n" + "\n".join(table_text) + "\n")
            
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return ""

def get_file_text(file_path: str) -> str:
    """
    Get text content from a file based on its extension.
    Supports PDF and DOCX formats.
    """
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        print(f"Unsupported file format: {file_path}")
        return ""

def extract_rfp_metadata(text: str) -> Dict[str, Any]:
    """
    Extract key metadata from RFP document text.
    Looks for common RFP elements like deadlines, client info, etc.
    """
    metadata = {
        "deadlines": [],
        "client_info": {},
        "submission_requirements": []
    }
    
    # Extract submission deadlines using regex patterns
    deadline_patterns = [
        r"due\s+(?:date|by)?\s*:?\s*((?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4})",
        r"deadline\s*:?\s*((?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4})",
        r"submit(?:ted)?\s+by\s*:?\s*((?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4})",
        r"due\s+(?:date|by)?\s*:?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
        r"deadline\s*:?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    ]
    
    for pattern in deadline_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            metadata["deadlines"].extend(matches)
    
    # Extract client information
    client_patterns = [
        r"(?:client|customer|company)\s+name\s*:?\s*([A-Za-z0-9\s\.,]+)(?:\n|$)",
        r"(?:prepared for|submitted to)\s*:?\s*([A-Za-z0-9\s\.,]+)(?:\n|$)"
    ]
    
    for pattern in client_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            metadata["client_info"]["name"] = match.group(1).strip()
            break
    
    # Extract submission requirements
    req_patterns = [
        r"(?:submission|format|deliver\w*)\s+requirements\s*:([^:]+)(?:\n\n|\n[A-Z])",
        r"proposals\s+must\s+be\s+submitted\s+([^\.]+)\."
    ]
    
    for pattern in req_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            metadata["submission_requirements"].extend([m.strip() for m in matches])
    
    return metadata

def extract_questions(text: str) -> List[Dict[str, str]]:
    """
    Extract questions from RFP document text.
    Identifies explicit questions and implied requirements.
    """
    questions = []
    
    # Extract explicit questions (ending with question mark)
    explicit_questions = re.findall(r"([A-Za-z0-9\s\.,;:'\"\(\)\-\/]{10,}[?])", text)
    for i, q in enumerate(explicit_questions):
        questions.append({
            "id": f"q{i+1}",
            "text": q.strip(),
            "type": "explicit",
            "section": "General"
        })
    
    # Extract numbered items that might be implicit questions/requirements
    numbered_items = re.findall(r"(?:\n|\r\n)(\d+\.?\d*\.?\s+[A-Za-z0-9\s\.,;:'\"\(\)\-\/]{10,})(?:\n|\r\n)", text)
    for i, item in enumerate(numbered_items):
        # Skip if it's too short or likely not a question/requirement
        if len(item.strip()) < 20:
            continue
            
        questions.append({
            "id": f"r{i+1}",
            "text": item.strip(),
            "type": "requirement",
            "section": "Requirements"
        })
    
    # Extract section headers to categorize questions
    sections = re.findall(r"(?:\n|\r\n)((?:[A-Z][A-Z\s]+:|(?:\d+\.){1,2}\s+[A-Z][A-Za-z\s]+))", text)
    
    # Assign sections to questions based on their position in the text
    current_section = "General"
    for section in sections:
        section_pos = text.find(section)
        section_name = section.strip().rstrip(':')
        
        # Update questions that appear after this section
        for q in questions:
            q_pos = text.find(q["text"])
            if q_pos > section_pos:
                q["section"] = section_name
                
        current_section = section_name
    
    return questions

def create_document_chunks(text: str, chunk_size: int = 2048, chunk_overlap: int = 300) -> List[Document]:
    """
    Split document text into chunks for processing.
    Uses intelligent splitting to preserve context and structure.
    """
    # First try to split by sections
    section_pattern = r"(?:\n|\r\n)(?:\d+\.\d+\s+|\d+\.\s+|[A-Z][A-Z\s]+:)"
    sections = re.split(section_pattern, text)
    
    chunks = []
    current_chunk = ""
    current_metadata = {"source": "document"}
    
    # Process each section
    for i, section in enumerate(sections):
        if not section.strip():
            continue
            
        # If section fits in chunk, add it
        if len(current_chunk) + len(section) <= chunk_size:
            current_chunk += section
        else:
            # If current chunk has content, save it
            if current_chunk:
                chunks.append(Document(
                    page_content=current_chunk,
                    metadata=current_metadata
                ))
            
            # Start new chunk
            if len(section) <= chunk_size:
                current_chunk = section
            else:
                # Split large sections into smaller chunks
                words = section.split()
                current_chunk = ""
                
                for word in words:
                    if len(current_chunk) + len(word) + 1 <= chunk_size:
                        current_chunk += word + " "
                    else:
                        chunks.append(Document(
                            page_content=current_chunk,
                            metadata=current_metadata
                        ))
                        
                        # Start new chunk with overlap
                        overlap_words = current_chunk.split()[-chunk_overlap:]
                        current_chunk = " ".join(overlap_words) + " " + word + " "
    
    # Add final chunk if it has content
    if current_chunk:
        chunks.append(Document(
            page_content=current_chunk,
            metadata=current_metadata
        ))
    
    return chunks
